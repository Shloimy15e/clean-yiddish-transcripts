"""
Document processor for handling Word documents.

Uses the plugin-based architecture for cleaning and output writers.
"""

from typing import Dict, Any, Optional, List
from pathlib import Path

from docx import Document

from cleaner import TranscriptCleaner
from registry import WriterRegistry
from converter import convert_doc_to_docx

# Import writers to register them (side effect: registers them)
import writers.docx_writer  # noqa: F401
import writers.txt_writer  # noqa: F401

# Default processors to use when none specified
# Note: 'brackets_inline' replaces 'regex' for smarter bracket handling
# Note: 'parentheses_notes' is NOT included by default since most parens are spoken
DEFAULT_PROCESSORS = ['special_chars', 'seif_marker', 'title_style', 'brackets_inline', 'whitespace']


class DocumentProcessor:
    """Processes Word documents and extracts text."""
    
    # XML namespaces used in Word documents
    WORD_NAMESPACES = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'wps': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape',
        'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
        'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
        'v': 'urn:schemas-microsoft-com:vml',
    }
    
    def __init__(self):
        self.cleaner = TranscriptCleaner()
    
    def _extract_textbox_text_from_paragraph(self, para) -> Optional[str]:
        """
        Extract text from text boxes/frames embedded in a paragraph.
        
        Text boxes are used for drop caps, special formatting, etc.
        They appear as w:txbxContent or within drawing elements.
        
        Returns:
            Optional[str]: Text from text boxes, or None if no text boxes found
        """
        try:
            para_xml = para._element
            
            # Try multiple XPath patterns for different text box formats
            textbox_patterns = [
                # Standard Word 2010+ text boxes
                './/wps:txbx//w:t',
                # Older VML text boxes
                './/v:textbox//w:t',
                # Alternative pattern
                './/w:txbxContent//w:t',
                # Drawing text boxes
                './/mc:AlternateContent//wps:txbx//w:t',
            ]
            
            textbox_texts = []
            for pattern in textbox_patterns:
                try:
                    texts = para_xml.xpath(pattern, namespaces=self.WORD_NAMESPACES)
                    for t in texts:
                        if t.text:
                            textbox_texts.append(t.text)
                except:
                    continue
            
            if textbox_texts:
                return ''.join(textbox_texts)
            
            return None
        except Exception:
            return None
    
    def _extract_all_textboxes_from_doc(self, doc) -> List[str]:
        """
        Extract all text box contents from the entire document.
        
        Returns:
            List[str]: List of text box contents
        """
        try:
            xml_content = doc._element.getroottree()
            
            textbox_patterns = [
                '//wps:txbx//w:t',
                '//v:textbox//w:t', 
                '//w:txbxContent//w:t',
            ]
            
            all_textbox_texts = []
            for pattern in textbox_patterns:
                try:
                    texts = xml_content.xpath(pattern, namespaces=self.WORD_NAMESPACES)
                    current_box = []
                    for t in texts:
                        if t.text:
                            current_box.append(t.text)
                    if current_box:
                        all_textbox_texts.append(''.join(current_box))
                except:
                    continue
            
            return all_textbox_texts
        except Exception:
            return []
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """
        Extract text from a .docx file.
        
        Args:
            file_path: Path to the .docx file
            
        Returns:
            str: Extracted text
        """
        try:
            doc = Document(file_path)
            paragraphs = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            return '\n'.join(paragraphs)
        except Exception as e:
            raise Exception(f"Error reading document: {str(e)}")
    
    def extract_paragraphs_with_metadata(self, file_path: str) -> tuple:
        """
        Extract paragraphs with their metadata (style, font size, etc.).
        
        Handles text boxes/frames that may contain content that should be
        merged with adjacent paragraphs (e.g., drop caps).
        
        Args:
            file_path: Path to the .docx file
            
        Returns:
            tuple: (full_text, list of paragraph metadata dicts)
        """
        try:
            doc = Document(file_path)
            paragraphs_meta = []
            all_font_sizes = []
            
            doc_default_size = 12
            try:
                normal_style = doc.styles['Normal']
                if normal_style.font and normal_style.font.size:
                    doc_default_size = normal_style.font.size.pt
            except KeyError:
                pass
            
            # Extract all text boxes from document for reference
            all_textboxes = self._extract_all_textboxes_from_doc(doc)
            
            # Track which textbox content has been used
            used_textbox_content = set()
            
            current_pos = 0
            for para in doc.paragraphs:
                # Get base paragraph text
                para_text = para.text
                
                # Check for text boxes embedded in this paragraph
                textbox_text = self._extract_textbox_text_from_paragraph(para)
                
                # If there's text box content, prepend it to the paragraph
                if textbox_text and textbox_text.strip():
                    # Only add if it's not already at the start of the paragraph
                    if not para_text.startswith(textbox_text):
                        para_text = textbox_text + para_text
                    used_textbox_content.add(textbox_text)
                
                if not para_text.strip():
                    continue
                
                style_name = para.style.name if para.style else None
                is_heading_style = style_name and ('heading' in style_name.lower() or 'title' in style_name.lower())
                
                font_size = self._get_paragraph_font_size(para, doc_default_size)
                if font_size:
                    all_font_sizes.append(font_size)
                
                is_bold = self._is_paragraph_bold(para)
                
                # Extract runs with formatting
                runs = self._extract_runs(para)
                
                para_len = len(para_text)
                paragraphs_meta.append({
                    'text': para_text,
                    'original_text': para_text,  # Keep original for highlighting
                    'start_pos': current_pos,
                    'end_pos': current_pos + para_len,
                    'style_name': style_name,
                    'is_heading_style': is_heading_style,
                    'is_bold': is_bold,
                    'font_size': font_size,
                    'char_count': para_len,
                    'word_count': len(para_text.split()),
                    'runs': runs,  # Store run formatting for output
                    'had_textbox': textbox_text is not None,  # Track if we merged a textbox
                })
                # +1 for the newline separator
                current_pos += para_len + 1
            
            # Check for orphaned textbox content that wasn't merged
            # This handles cases where textbox is a separate paragraph
            self._merge_orphaned_textboxes(paragraphs_meta, all_textboxes, used_textbox_content)
            
            avg_font_size = sum(all_font_sizes) / len(all_font_sizes) if all_font_sizes else 12
            
            for meta in paragraphs_meta:
                meta['is_larger_than_normal'] = (
                    meta['font_size'] is not None and 
                    meta['font_size'] > avg_font_size * 1.2
                )
                meta['avg_font_size'] = avg_font_size
            
            full_text = '\n'.join(p['text'] for p in paragraphs_meta)
            return full_text, paragraphs_meta
            
        except Exception as e:
            raise Exception(f"Error reading document metadata: {str(e)}")
    
    def _merge_orphaned_textboxes(self, paragraphs_meta: List[Dict], 
                                   all_textboxes: List[str], 
                                   used_textbox_content: set) -> None:
        """
        Merge orphaned textbox content that appears as separate short paragraphs.
        
        This handles cases where a text box (like a drop cap) appears as its own
        paragraph rather than being embedded in the following paragraph.
        
        Args:
            paragraphs_meta: List of paragraph metadata to modify in-place
            all_textboxes: All textbox content found in document
            used_textbox_content: Set of textbox content already merged
        """
        if len(paragraphs_meta) < 2:
            return
        
        # Look for very short paragraphs (1-2 words) that might be orphaned textbox content
        indices_to_remove = []
        
        for i in range(len(paragraphs_meta) - 1):
            para = paragraphs_meta[i]
            next_para = paragraphs_meta[i + 1]
            
            # Check if this paragraph looks like orphaned textbox content:
            # - Very short (1-3 words, less than 20 chars)
            # - Followed by a longer paragraph
            # - Not a heading
            para_text = para['text'].strip()
            if (para['word_count'] <= 3 and 
                len(para_text) < 20 and
                not para.get('is_heading_style') and
                next_para['word_count'] > 3):
                
                # Check if this text appears in textbox list or matches textbox pattern
                # Only merge if it's actually textbox content that hasn't been used yet
                matched_textbox = None
                is_likely_textbox = False
                
                # Check for exact match first
                if para_text in all_textboxes and para_text not in used_textbox_content:
                    is_likely_textbox = True
                    matched_textbox = para_text
                else:
                    # Check for partial match - para_text is the start of a textbox
                    # This handles drop caps where "T" might be the first char of a textbox
                    for tb in all_textboxes:
                        if tb not in used_textbox_content and tb.startswith(para_text):
                            is_likely_textbox = True
                            matched_textbox = tb
                            break
                
                if is_likely_textbox:
                    # Merge with next paragraph
                    merged_text = para_text + ' ' + next_para['text']
                    next_para['text'] = merged_text
                    next_para['original_text'] = merged_text
                    next_para['word_count'] = len(merged_text.split())
                    next_para['char_count'] = len(merged_text)
                    next_para['had_textbox_merged'] = True
                    indices_to_remove.append(i)
                    # Mark both the paragraph text and the matched textbox as used
                    used_textbox_content.add(para_text)
                    if matched_textbox:
                        used_textbox_content.add(matched_textbox)
        
        # Remove merged paragraphs in reverse order to maintain indices
        for i in reversed(indices_to_remove):
            paragraphs_meta.pop(i)
        
        # Recalculate positions after merging
        current_pos = 0
        for meta in paragraphs_meta:
            para_len = len(meta['text'])
            meta['start_pos'] = current_pos
            meta['end_pos'] = current_pos + para_len
            current_pos += para_len + 1

    def _extract_runs(self, para) -> List[Dict[str, Any]]:
        """Extract runs with their formatting from a paragraph."""
        runs = []
        for run in para.runs:
            run_data = {
                'text': run.text,
                'style': {
                    'bold': run.font.bold,
                    'italic': run.font.italic,
                    'underline': run.font.underline is not None and run.font.underline,
                    'font_size': run.font.size.pt if run.font.size else None,
                    'font_name': run.font.name,
                    'strike': run.font.strike,
                    'superscript': run.font.superscript,
                    'subscript': run.font.subscript,
                }
            }
            # Extract color if present
            if run.font.color and run.font.color.rgb:
                try:
                    rgb_obj = run.font.color.rgb
                    hex_color = str(rgb_obj)
                    if len(hex_color) == 6:
                        r = int(hex_color[0:2], 16)
                        g = int(hex_color[2:4], 16)
                        b = int(hex_color[4:6], 16)
                        run_data['style']['color_rgb'] = (r, g, b)
                except Exception:
                    pass
            
            runs.append(run_data)
        
        return runs
    
    def _is_paragraph_bold(self, para) -> bool:
        """Check if the entire paragraph is bold."""
        if not para.runs:
            return False
        for run in para.runs:
            if run.text.strip() and not run.bold:
                return False
        return len(para.runs) > 0 and any(r.text.strip() for r in para.runs)
    
    def _get_paragraph_font_size(self, para, doc_default_size=None) -> Optional[float]:
        """Get the font size of a paragraph (from first run, style, or default)."""
        for run in para.runs:
            if run.font.size:
                return run.font.size.pt
        
        if para.style and para.style.font and para.style.font.size:
            return para.style.font.size.pt
        
        style = para.style
        while style:
            if style.font and style.font.size:
                return style.font.size.pt
            style = style.base_style
        
        return doc_default_size
    
    def process_document(self, file_path: str, filename: str,
                         processors: Optional[List[str]] = None) -> Dict[str, Any]:
        """
        Process a document: extract text, clean it, and return results.

        Args:
            file_path: Path to the document file
            filename: Original filename
            processors: List of processor names to apply (default: DEFAULT_PROCESSORS)

        Returns:
            dict: Processing results including original, cleaned, removed items, and stats
        """
        # Convert .doc to .docx if necessary
        file_path = Path(file_path)
        temp_docx_path = None

        if file_path.suffix.lower() == '.doc':
            temp_docx_path = convert_doc_to_docx(str(file_path))
            file_path = Path(temp_docx_path)

        try:
            original_text, paragraphs_meta = self.extract_paragraphs_with_metadata(str(file_path))
        finally:
            # Clean up temporary .docx file if we created one
            if temp_docx_path and Path(temp_docx_path).exists():
                Path(temp_docx_path).unlink()
        
        context = {
            'paragraphs': paragraphs_meta,
        }
        
        # Use provided processors or default set
        processor_list = processors if processors else DEFAULT_PROCESSORS
        
        cleaned_text, removed_items = self.cleaner.clean_with_processors(
            original_text, processor_list, context
        )
        
        stats = self.cleaner.get_statistics(original_text, cleaned_text)
        
        return {
            'filename': filename,
            'original_text': original_text,
            'cleaned_text': cleaned_text,
            'removed_items': removed_items,
            'statistics': stats,
            'processors': processor_list,
            'context': context,  # Include context for writers
            'success': True
        }
    
    def save_cleaned_document(self, cleaned_text: str, output_path: str,
                              format_name: str = 'docx',
                              context: Optional[Dict[str, Any]] = None) -> str:
        """
        Save cleaned text to a document using the specified writer.
        
        Args:
            cleaned_text: The cleaned text to save
            output_path: Path where to save the document
            format_name: Output format ('docx', 'txt', etc.)
            context: Optional context with paragraph metadata for formatting
            
        Returns:
            str: The output path
        """
        writer = WriterRegistry.get_writer(format_name)
        if not writer:
            # Fallback to docx
            writer = WriterRegistry.get_writer('docx')
        
        if writer:
            writer.write(cleaned_text, output_path, context)
        else:
            # Ultimate fallback - basic docx
            from docx import Document as DocxDoc
            doc = DocxDoc()
            for para_text in cleaned_text.split('\n'):
                if para_text.strip():
                    doc.add_paragraph(para_text)
            doc.save(output_path)
        
        return output_path
    
    def get_cleaned_bytes(self, cleaned_text: str, format_name: str = 'docx',
                          context: Optional[Dict[str, Any]] = None) -> bytes:
        """
        Get cleaned document as bytes for download.
        
        Args:
            cleaned_text: The cleaned text
            format_name: Output format ('docx', 'txt', etc.)
            context: Optional context with paragraph metadata
            
        Returns:
            bytes: The document content as bytes
        """
        writer = WriterRegistry.get_writer(format_name)
        if not writer:
            writer = WriterRegistry.get_writer('docx')
        
        if writer:
            return writer.write_to_bytes(cleaned_text, context)
        else:
            # Fallback
            return cleaned_text.encode('utf-8')
    
    def get_available_formats(self) -> Dict[str, Dict[str, str]]:
        """Get available output formats."""
        return WriterRegistry.get_formats()
