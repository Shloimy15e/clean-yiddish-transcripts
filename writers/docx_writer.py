"""
Writer for Microsoft Word .docx files with RTL and formatting preservation.
"""

from pathlib import Path
from typing import Dict, Any, Optional, Union, List
from io import BytesIO

from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from registry import WriterRegistry
from writers.base import OutputWriter
from utils import sanitize_xml_text


@WriterRegistry.register
class DocxWriter(OutputWriter):
    """Writer for Microsoft Word .docx files with RTL support."""
    
    name = "docx"
    title = "Word Document"
    description = "Microsoft Word document with RTL support and preserved formatting"
    extension = ".docx"
    mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    
    def write(self, text: str, output: Union[str, Path, BytesIO],
              context: Optional[Dict[str, Any]] = None) -> None:
        """
        Write the cleaned text to a .docx file.
        
        Args:
            text: The cleaned text to write
            output: Output path (str/Path) or BytesIO stream
            context: Optional context with 'paragraphs' list containing paragraph metadata
        """
        doc = self._create_document(text, context)
        
        if isinstance(output, BytesIO):
            doc.save(output)
        else:
            output_path = Path(output)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(output_path))
    
    def write_to_bytes(self, text: str,
                       context: Optional[Dict[str, Any]] = None) -> bytes:
        """
        Write the cleaned text and return as bytes.
        
        Args:
            text: The cleaned text to write
            context: Optional context with paragraph metadata
            
        Returns:
            bytes: The .docx file content as bytes
        """
        buffer = BytesIO()
        self.write(text, buffer, context)
        buffer.seek(0)
        return buffer.read()
    
    def _create_document(self, text: str, 
                         context: Optional[Dict[str, Any]] = None) -> DocxDocument:
        """Create a Word document from the text."""
        doc = DocxDocument()
        
        # Configure default styles
        self._configure_styles(doc)
        
        # Get paragraph metadata if available
        paragraphs_meta = None
        if context and 'paragraphs' in context:
            paragraphs_meta = context['paragraphs']
        
        # Split text into paragraphs
        paragraphs = text.split('\n')
        
        for i, para_text in enumerate(paragraphs):
            para_text = para_text.strip()
            if not para_text:
                continue
            
            # Sanitize for XML
            para_text = sanitize_xml_text(para_text)
            
            # Create paragraph with RTL
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.right_to_left = True
            
            # Check if we have metadata for this paragraph
            meta = None
            if paragraphs_meta and i < len(paragraphs_meta):
                meta = paragraphs_meta[i]
            
            # Add runs with formatting
            if meta and 'runs' in meta:
                self._add_formatted_runs(p, meta['runs'])
            else:
                # Simple text run
                run = p.add_run(para_text)
                run.font.size = Pt(12)
        
        return doc
    
    def _configure_styles(self, doc: DocxDocument) -> None:
        """Configure document styles for Hebrew/RTL text."""
        styles = doc.styles
        
        try:
            normal = styles['Normal']
            normal.font.size = Pt(12)
            normal.paragraph_format.space_after = Pt(0)
            normal.paragraph_format.line_spacing = 1.15
        except KeyError:
            pass
    
    def _add_formatted_runs(self, paragraph, runs: List[Dict[str, Any]]) -> None:
        """Add formatted text runs to a paragraph."""
        for run_data in runs:
            text = run_data.get('text', '')
            if not text:
                continue
            
            text = sanitize_xml_text(text)
            run = paragraph.add_run(text)
            
            style = run_data.get('style', {})
            
            if style.get('bold'):
                run.font.bold = True
            if style.get('italic'):
                run.font.italic = True
            if style.get('underline'):
                run.font.underline = True
            if style.get('font_size'):
                run.font.size = Pt(style['font_size'])
            if style.get('font_name'):
                run.font.name = style['font_name']
            if style.get('color_rgb'):
                r, g, b = style['color_rgb']
                run.font.color.rgb = RGBColor(r, g, b)
            if style.get('strike'):
                run.font.strike = True
            if style.get('superscript'):
                run.font.superscript = True
            if style.get('subscript'):
                run.font.subscript = True
